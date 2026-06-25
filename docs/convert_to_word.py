"""
VattalettuX — Markdown to Word Document Converter

Converts all chapter markdown files into a single formatted Word document
for final year MCA project submission.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Configuration ─────────────────────────────────────────────────────
CHAPTERS_DIR = os.path.join(os.path.dirname(__file__), "chapters")
OUTPUT_FILE = os.path.join(os.path.dirname(__file__), "VattalettuX_Project_Report.docx")

CHAPTER_FILES = [
    "00_front_matter.md",
    "01_abstract.md",
    "02_introduction.md",
    "03_literature_survey.md",
    "04_system_analysis.md",
    "05_system_design.md",
    "06_implementation.md",
    "07_testing_results.md",
    "08_conclusion.md",
    "09_references.md",
    "10_appendices.md",
]

FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = 12
FONT_SIZE_H1 = 16
FONT_SIZE_H2 = 14
FONT_SIZE_H3 = 13
FONT_SIZE_H4 = 12
LINE_SPACING = 1.5


# ── Helper: set cell shading ──────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


# ── Helper: format run ───────────────────────────────────────────────
def format_run(run, font_name=FONT_NAME, font_size=FONT_SIZE_BODY,
               bold=False, italic=False, color=None):
    """Apply formatting to a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Force Times New Roman for East Asian text too
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)


# ── Helper: add formatted paragraph ──────────────────────────────────
def add_paragraph(doc, text="", style=None, font_size=FONT_SIZE_BODY,
                  bold=False, italic=False, alignment=None,
                  space_before=0, space_after=6, color=None,
                  first_line_indent=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = LINE_SPACING
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)

    if text:
        run = p.add_run(text)
        format_run(run, font_size=font_size, bold=bold, italic=italic, color=color)
    return p


# ── Helper: parse inline formatting (bold, italic, code, links) ──────
def add_inline_formatted_text(paragraph, text, base_size=FONT_SIZE_BODY):
    """Parse markdown inline formatting and add runs to a paragraph."""
    if not text:
        return

    # Split on inline patterns: **bold**, *italic*, `code`, [text](url)
    # Process from left to right
    i = 0
    while i < len(text):
        # Bold: **text**
        bold_match = re.match(r'\*\*(.+?)\*\*', text[i:])
        if bold_match:
            run = paragraph.add_run(bold_match.group(1))
            format_run(run, font_size=base_size, bold=True)
            i += bold_match.end()
            continue

        # Italic: *text*
        italic_match = re.match(r'\*(.+?)\*', text[i:])
        if italic_match:
            run = paragraph.add_run(italic_match.group(1))
            format_run(run, font_size=base_size, italic=True)
            i += italic_match.end()
            continue

        # Inline code: `text`
        code_match = re.match(r'`(.+?)`', text[i:])
        if code_match:
            run = paragraph.add_run(code_match.group(1))
            format_run(run, font_name="Consolas", font_size=base_size - 1,
                      color=(128, 0, 0))
            i += code_match.end()
            continue

        # Link: [text](url)
        link_match = re.match(r'\[(.+?)\]\((.+?)\)', text[i:])
        if link_match:
            run = paragraph.add_run(link_match.group(1))
            format_run(run, font_size=base_size, color=(0, 0, 200))
            i += link_match.end()
            continue

        # Regular character
        # Collect plain text until next special character
        plain_match = re.match(r'([^*`\[]+)', text[i:])
        if plain_match:
            run = paragraph.add_run(plain_match.group(1))
            format_run(run, font_size=base_size)
            i += plain_match.end()
        else:
            run = paragraph.add_run(text[i])
            format_run(run, font_size=base_size)
            i += 1


# ── Helper: add a table from markdown ─────────────────────────────────
def add_table_from_rows(doc, header_row, data_rows):
    """Create a formatted Word table from parsed markdown table data."""
    num_cols = len(header_row)
    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for j, cell_text in enumerate(header_row):
        cell = hdr.cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_text.strip())
        format_run(run, font_size=11, bold=True)
        set_cell_shading(cell, "2E4057")
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_data in data_rows:
        row = table.add_row()
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                # Remove markdown bold from cell text for cleaner display
                clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text.strip())
                clean_text = re.sub(r'`(.+?)`', r'\1', clean_text)
                run = p.add_run(clean_text)
                format_run(run, font_size=10)

    # Add spacing after table
    add_paragraph(doc, "", space_after=6)
    return table


# ── Helper: add code block ───────────────────────────────────────────
def add_code_block(doc, code_lines):
    """Add a formatted code block."""
    code_text = "\n".join(code_lines)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    pf.left_indent = Cm(1)

    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)

    # Set East Asian font too
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), "Consolas")
    rFonts.set(qn('w:hAnsi'), "Consolas")
    rFonts.set(qn('w:cs'), "Consolas")


# ── Helper: add ASCII art/diagram block ──────────────────────────────
def add_diagram_block(doc, lines):
    """Add a text-based diagram block with monospace font."""
    text = "\n".join(lines)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    pf.left_indent = Cm(0.5)

    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(40, 40, 40)

    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), "Consolas")
    rFonts.set(qn('w:hAnsi'), "Consolas")
    rFonts.set(qn('w:cs'), "Consolas")


# ── Main: parse markdown and build Word doc ──────────────────────────
def parse_markdown_line(line):
    """Determine the type of a markdown line."""
    stripped = line.strip()

    if stripped.startswith("# "):
        return ("h1", stripped[2:].strip())
    elif stripped.startswith("## "):
        return ("h2", stripped[3:].strip())
    elif stripped.startswith("### "):
        return ("h3", stripped[4:].strip())
    elif stripped.startswith("#### "):
        return ("h4", stripped[5:].strip())
    elif stripped == "---":
        return ("hr", "")
    elif stripped.startswith("| ") and "|" in stripped[1:]:
        return ("table_row", stripped)
    elif re.match(r'^\d+\.\s', stripped):
        return ("ordered_list", re.sub(r'^\d+\.\s', '', stripped))
    elif stripped.startswith("- ") or stripped.startswith("* "):
        return ("unordered_list", stripped[2:])
    elif stripped.startswith("> "):
        return ("blockquote", stripped[2:])
    elif stripped.startswith("```"):
        return ("code_fence", stripped[3:].strip())
    elif stripped.startswith("*Figure") or stripped.startswith("*Note") or stripped.startswith("*["):
        return ("italic_note", stripped.strip("*").strip())
    elif stripped == "":
        return ("empty", "")
    else:
        return ("text", stripped)


def is_ascii_art(lines):
    """Check if code block lines look like ASCII art/diagram."""
    art_chars = set("┌┐└┘│─├┤┬┴┼╌╎═║╔╗╚╝╠╣╦╩╬▲▼◄►←→↑↓│┌┐└┘├┤┬┴┼▶◀")
    for line in lines:
        for ch in line:
            if ch in art_chars:
                return True
    return False


def convert_chapter(doc, filepath, add_page_break=True):
    """Convert a single markdown chapter file to Word content."""
    print(f"  Processing: {os.path.basename(filepath)}")

    with open(filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()

    if add_page_break:
        doc.add_page_break()

    in_code_block = False
    code_lines = []
    code_lang = ""
    in_table = False
    table_rows = []
    table_header = None
    skip_separator = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n").rstrip("\r")
        line_type, content = parse_markdown_line(line)

        # ── Code block handling ──
        if line_type == "code_fence":
            if not in_code_block:
                in_code_block = True
                code_lines = []
                code_lang = content
            else:
                # End code block
                in_code_block = False
                if code_lines:
                    if is_ascii_art(code_lines):
                        add_diagram_block(doc, code_lines)
                    else:
                        add_code_block(doc, code_lines)
                code_lines = []
                code_lang = ""
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Table handling ──
        if line_type == "table_row":
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if not in_table:
                in_table = True
                table_header = cells
                skip_separator = True
            elif skip_separator:
                # This is the separator row (|---|---|)
                skip_separator = False
            else:
                table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table
            if table_header and table_rows:
                add_table_from_rows(doc, table_header, table_rows)
            elif table_header:
                add_table_from_rows(doc, table_header, [])
            in_table = False
            table_rows = []
            table_header = None
            skip_separator = False
            # Don't increment i — process current line normally
            continue

        # ── Regular content ──
        if line_type == "h1":
            p = add_paragraph(doc, content, font_size=FONT_SIZE_H1,
                            bold=True, space_before=18, space_after=12,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        elif line_type == "h2":
            p = add_paragraph(doc, "", font_size=FONT_SIZE_H2,
                            bold=True, space_before=14, space_after=8)
            add_inline_formatted_text(p, content, base_size=FONT_SIZE_H2)
            for run in p.runs:
                run.font.bold = True
        elif line_type == "h3":
            p = add_paragraph(doc, "", font_size=FONT_SIZE_H3,
                            bold=True, space_before=10, space_after=6)
            add_inline_formatted_text(p, content, base_size=FONT_SIZE_H3)
            for run in p.runs:
                run.font.bold = True
        elif line_type == "h4":
            p = add_paragraph(doc, "", font_size=FONT_SIZE_H4,
                            bold=True, italic=True, space_before=8, space_after=4)
            add_inline_formatted_text(p, content, base_size=FONT_SIZE_H4)
            for run in p.runs:
                run.font.bold = True
                run.font.italic = True
        elif line_type == "hr":
            # Just add some spacing
            add_paragraph(doc, "", space_before=4, space_after=4)
        elif line_type == "ordered_list":
            p = add_paragraph(doc, "", space_before=2, space_after=2)
            p.paragraph_format.left_indent = Cm(1.5)
            # Find the number
            num_match = re.match(r'^(\d+)\.\s', line.strip())
            num_str = num_match.group(1) + ". " if num_match else "• "
            run = p.add_run(num_str)
            format_run(run, font_size=FONT_SIZE_BODY, bold=True)
            add_inline_formatted_text(p, content, base_size=FONT_SIZE_BODY)
        elif line_type == "unordered_list":
            p = add_paragraph(doc, "", space_before=2, space_after=2)
            p.paragraph_format.left_indent = Cm(1.5)
            run = p.add_run("• ")
            format_run(run, font_size=FONT_SIZE_BODY)
            add_inline_formatted_text(p, content, base_size=FONT_SIZE_BODY)
        elif line_type == "blockquote":
            p = add_paragraph(doc, "", space_before=4, space_after=4)
            p.paragraph_format.left_indent = Cm(1.5)
            clean = content.lstrip("> ").strip()
            if clean.startswith("[!NOTE]") or clean.startswith("[!TIP]"):
                clean = clean.split("]", 1)[-1].strip()
            add_inline_formatted_text(p, clean, base_size=FONT_SIZE_BODY)
            for run in p.runs:
                run.font.italic = True
        elif line_type == "italic_note":
            p = add_paragraph(doc, "", space_before=2, space_after=4)
            add_inline_formatted_text(p, content, base_size=FONT_SIZE_BODY)
            for run in p.runs:
                run.font.italic = True
        elif line_type == "text":
            p = add_paragraph(doc, "", space_before=0, space_after=6)
            add_inline_formatted_text(p, content, base_size=FONT_SIZE_BODY)
        elif line_type == "empty":
            pass  # Skip empty lines (spacing handled by paragraph formatting)

        i += 1

    # Flush any remaining table
    if in_table and table_header:
        add_table_from_rows(doc, table_header, table_rows if table_rows else [])


def setup_document():
    """Create and configure the base Word document."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE_BODY)
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(6)

    # Set East Asian font
    rFonts = style.element.find(qn('w:rPr'))
    if rFonts is None:
        rPr_elem = parse_xml(f'<w:rPr {nsdecls("w")}><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}" w:cs="{FONT_NAME}"/></w:rPr>')
        style.element.append(rPr_elem)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(2.54)

    return doc


def main():
    """Main conversion function."""
    print("=" * 60)
    print("VattalettuX — Markdown to Word Converter")
    print("=" * 60)

    doc = setup_document()

    for idx, chapter_file in enumerate(CHAPTER_FILES):
        filepath = os.path.join(CHAPTERS_DIR, chapter_file)
        if not os.path.exists(filepath):
            print(f"  WARNING: {chapter_file} not found, skipping")
            continue
        convert_chapter(doc, filepath, add_page_break=(idx > 0))

    # Add page numbers in footer
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page number field
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)

        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)

        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)

        format_run(run, font_size=10)
        format_run(run2, font_size=10)
        format_run(run3, font_size=10)

    doc.save(OUTPUT_FILE)
    print(f"\n{'=' * 60}")
    print(f"SUCCESS! Document saved to:")
    print(f"  {OUTPUT_FILE}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
