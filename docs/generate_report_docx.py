"""
Convert expanded MCA Project Report from Markdown to Word Document (.docx)
Reads the combined markdown file and generates a properly formatted Word document.
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Setup
MD_FILE = r"f:\final mca project\VattalettuX\docs\MCA_Project_Report_BHC_April2026.md"
OUT_FILE = r"f:\final mca project\VattalettuX\docs\VatteluttuX_Report_Final.docx"
IMG_DIR = r"C:\Users\Asus\.gemini\antigravity\brain\c39d4a0b-fd6d-447d-a543-9be6e133db06"

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)

# Helper functions
def set_run_font(run, size=12, bold=False, italic=False, name='Times New Roman'):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_formatted_para(text, size=12, bold=False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p

def add_heading_tnr(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_table_from_rows(header_row, data_rows):
    if not header_row:
        return
    ncols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = h.strip()
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
    for ri, row in enumerate(data_rows):
        for ci in range(min(ncols, len(row))):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = row[ci].strip()
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(10)
    doc.add_paragraph()

def add_code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(8)

def add_image_if_exists(img_path, caption=""):
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.0))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

# Read markdown
with open(MD_FILE, 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')
i = 0
in_code_block = False
code_lines = []
in_table = False
table_rows = []
list_counter = 0

while i < len(lines):
    line = lines[i]
    
    # Code block start/end
    if line.strip().startswith('```'):
        if in_code_block:
            add_code_block(code_lines)
            code_lines = []
            in_code_block = False
        else:
            in_code_block = True
            code_lines = []
        i += 1
        continue
    
    if in_code_block:
        code_lines.append(line)
        i += 1
        continue
    
    # Skip horizontal rules
    if line.strip() == '---':
        i += 1
        continue
    
    # Table detection
    if '|' in line and line.strip().startswith('|'):
        stripped = line.strip()
        cells = [c.strip() for c in stripped.split('|')[1:-1]]
        
        # Check if separator row
        if all(c.replace('-', '').replace(':', '') == '' for c in cells):
            i += 1
            continue
        
        if not in_table:
            in_table = True
            table_rows = [cells]
        else:
            table_rows.append(cells)
        
        # Check if next line continues table
        if i + 1 < len(lines) and '|' in lines[i + 1] and lines[i + 1].strip().startswith('|'):
            i += 1
            continue
        else:
            # End of table, render it
            if len(table_rows) > 1:
                add_table_from_rows(table_rows[0], table_rows[1:])
            elif len(table_rows) == 1:
                add_table_from_rows(table_rows[0], [])
            in_table = False
            table_rows = []
            i += 1
            continue
    
    # Image
    img_match = re.match(r'!\[(.+?)\]\((.+?)\)', line.strip())
    if img_match:
        caption = img_match.group(1)
        img_path = img_match.group(2)
        add_image_if_exists(img_path, caption)
        i += 1
        continue
    
    # Headings
    if line.startswith('# ') and not line.startswith('## '):
        text = line[2:].strip()
        # Page break before major chapters (but not for very first)
        if doc.paragraphs and ('CHAPTER' in text.upper() or 'APPENDIX' in text.upper() or 'CONTENTS' in text.upper()):
            doc.add_page_break()
        add_heading_tnr(text, level=1)
        list_counter = 0
        i += 1
        continue
    
    if line.startswith('## '):
        text = line[3:].strip()
        add_heading_tnr(text, level=2)
        list_counter = 0
        i += 1
        continue
    
    if line.startswith('### '):
        text = line[4:].strip()
        add_heading_tnr(text, level=3)
        list_counter = 0
        i += 1
        continue
    
    # Empty lines
    if line.strip() == '':
        i += 1
        continue
    
    # Numbered list items
    num_match = re.match(r'^(\d+)\.\s+(.+)', line.strip())
    if num_match:
        text = num_match.group(2)
        # Remove markdown bold
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        p = doc.add_paragraph(style='List Number')
        # Handle bold prefix with dash
        if ' â€” ' in text:
            parts = text.split(' â€” ', 1)
            run1 = p.add_run(parts[0])
            set_run_font(run1, bold=True)
            run2 = p.add_run(' â€” ' + parts[1])
            set_run_font(run2)
        elif ' â€“ ' in text:
            parts = text.split(' â€“ ', 1)
            run1 = p.add_run(parts[0])
            set_run_font(run1, bold=True)
            run2 = p.add_run(' â€“ ' + parts[1])
            set_run_font(run2)
        else:
            run = p.add_run(text)
            set_run_font(run)
        i += 1
        continue
    
    # Bullet list items
    if line.strip().startswith('- ') or line.strip().startswith('* '):
        text = line.strip()[2:]
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        set_run_font(run)
        i += 1
        continue
    
    # Regular paragraph
    text = line.strip()
    if text:
        # Remove markdown formatting for clean text
        clean = text
        p = doc.add_paragraph()
        
        # Parse bold sections
        parts = re.split(r'(\*\*.+?\*\*)', clean)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                set_run_font(run, bold=True)
            else:
                run = p.add_run(part)
                set_run_font(run)
    
    i += 1

# Save
doc.save(OUT_FILE)
print(f"\nWord document saved: {OUT_FILE}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
# Estimate pages
word_count = sum(len(p.text.split()) for p in doc.paragraphs)
est_pages = word_count / 300 + 6  # +6 for images/tables
print(f"Estimated word count: {word_count}")
print(f"Estimated pages: {est_pages:.0f}")




