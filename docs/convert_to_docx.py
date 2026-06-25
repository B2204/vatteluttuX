"""
Convert VattalettuX_Full_Documentation.md to a formatted Word document.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

DOCS_DIR = Path(__file__).parent
INPUT_FILE = DOCS_DIR / "VattalettuX_Full_Documentation.md"
OUTPUT_FILE = DOCS_DIR / "VattalettuX_Full_Documentation.docx"


def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    for level, (sz, sp_before) in {1: (16, 24), 2: (14, 18), 3: (13, 12)}.items():
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Times New Roman'
        h.font.size = Pt(sz)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.paragraph_format.space_before = Pt(sp_before)
        h.paragraph_format.space_after = Pt(8)
        if level == 1:
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        cs = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        cs = doc.styles['CodeBlock']
    cs.font.name = 'Consolas'
    cs.font.size = Pt(9)
    cs.paragraph_format.space_before = Pt(1)
    cs.paragraph_format.space_after = Pt(1)
    cs.paragraph_format.line_spacing = 1.0


def add_title_page(doc):
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("VatteluttuX")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.name = 'Times New Roman'

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Enhancing Epigraphical Research through\nDeep Learning-Based OCR and Modern Tamil Mapping")
    r.font.size = Pt(16)
    r.font.name = 'Times New Roman'

    doc.add_paragraph()
    doc.add_paragraph()

    l = doc.add_paragraph()
    l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = l.add_run("PROJECT DOCUMENTATION")
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.name = 'Times New Roman'

    for _ in range(6):
        doc.add_paragraph()

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = d.add_run("PG Department of Computer Applications")
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    doc.add_page_break()


def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and not all(c in '|-: ' for c in line):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)
    return rows


def add_table_to_doc(doc, rows):
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                clean = re.sub(r'`(.*?)`', r'\1', clean)
                run = p.add_run(clean)
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
                if i == 0:
                    run.bold = True
                    from docx.oxml import OxmlElement
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), 'D9E2F3')
                    shading_elm.set(qn('w:val'), 'clear')
                    cell._tc.get_or_add_tcPr().append(shading_elm)
    doc.add_paragraph()


def process_inline(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


def convert():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)

    setup_styles(doc)
    add_title_page(doc)

    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code = False
    code_lines = []
    in_table = False
    table_lines = []
    first_heading_skipped = False

    while i < len(lines):
        line = lines[i].rstrip('\r\n')

        # Skip very first title
        if not first_heading_skipped and line.startswith('# '):
            first_heading_skipped = True
            i += 1
            continue

        # Code block start
        if line.strip().startswith('```') and not in_code:
            if in_table and table_lines:
                add_table_to_doc(doc, parse_table(table_lines))
                table_lines = []
                in_table = False
            in_code = True
            code_lines = []
            i += 1
            continue

        # Code block end
        if line.strip() == '```' and in_code:
            for cl in code_lines:
                p = doc.add_paragraph(style='CodeBlock')
                run = p.add_run(cl)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            in_code = False
            code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            add_table_to_doc(doc, parse_table(table_lines))
            table_lines = []
            in_table = False

        # Skip ---
        if line.strip() == '---':
            i += 1
            continue

        # Skip empty
        if not line.strip():
            i += 1
            continue

        # Heading 1
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            if any(text.startswith(f'{n} ') for n in range(1, 10)) or text == 'APPENDIX':
                doc.add_page_break()
            doc.add_heading(text, level=1)
            i += 1
            continue

        # Heading 2
        if line.startswith('## ') and not line.startswith('### '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        # Heading 3
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            p = doc.add_paragraph(style='List Number')
            process_inline(p, m.group(2))
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        process_inline(p, line)
        i += 1

    # Flush remaining table
    if in_table and table_lines:
        add_table_to_doc(doc, parse_table(table_lines))

    doc.save(str(OUTPUT_FILE))
    print(f"SUCCESS! Word document saved to: {OUTPUT_FILE}")
    print(f"File size: {OUTPUT_FILE.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    convert()
